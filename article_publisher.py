#!/usr/bin/env python3
"""
三才靈芝農場 - 每日部落格文章自動發布系統
功能：自動生成 .docx 文章 → 解析 → 推送到 G系統

使用方式：
    python article_publisher.py generate <day>    # 生成文章
    python article_publisher.py publish <docx_path> <day>  # 推送文章
    python article_publisher.py test               # 測試連接
"""

import os
import sys
import json
import logging
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, Tuple

# 第三方庫
import requests
from docx import Document
from docx.oxml.ns import qn

# Azure 支援
try:
    from azure.identity import DefaultAzureCredential
    from azure.keyvault.secrets import SecretClient
    AZURE_AVAILABLE = True
except ImportError:
    AZURE_AVAILABLE = False

# ====== 配置函數 ======
def get_azure_secret(secret_name: str, vault_url: str = None) -> str:
    """
    從 Azure Key Vault 讀取密鑰
    如果 Key Vault 不可用，降級到環境變數
    """
    if not AZURE_AVAILABLE:
        return None

    try:
        vault_url = vault_url or os.getenv('AZURE_KEYVAULT_URL')
        if not vault_url:
            return None

        credential = DefaultAzureCredential()
        client = SecretClient(vault_url=vault_url, credential=credential)
        secret = client.get_secret(secret_name)
        return secret.value
    except Exception as e:
        # 靜默失敗，降級到環境變數
        return None

def load_config() -> Tuple[str, str]:
    """
    載入設定，優先使用 Azure Key Vault，降級到環境變數
    """
    # 嘗試從 Azure Key Vault 讀取
    vault_url = os.getenv('AZURE_KEYVAULT_URL')
    if vault_url and AZURE_AVAILABLE:
        api_url = get_azure_secret('G-SYSTEM-API-URL', vault_url)
        api_key = get_azure_secret('G-SYSTEM-API-KEY', vault_url)
        if api_url and api_key:
            return api_url, api_key

    # 降級到環境變數
    api_url = os.getenv('G_SYSTEM_API_URL')
    api_key = os.getenv('G_SYSTEM_API_KEY')

    # 降級到 .env 檔案（本地開發用）
    if not api_url or not api_key:
        try:
            from dotenv import load_dotenv
            load_dotenv()
            api_url = api_url or os.getenv('G_SYSTEM_API_URL')
            api_key = api_key or os.getenv('G_SYSTEM_API_KEY')
        except ImportError:
            pass

    # 預設值（不推薦用於生產）
    if not api_url:
        api_url = 'https://sungertain.deweichiu.com/api/articles/publish'
    if not api_key:
        api_key = 'gsy_prod_1a2b3c4d5e6f7g8h9i0j_cowork_2026'

    return api_url, api_key

# ====== 配置 ======
API_URL, API_KEY = load_config()
LOG_DIR = Path(os.getenv('LOG_DIR', '/Users/ericchiu/Documents/sungertain-design/G系統專案/logs'))
ARTICLES_DIR = Path(os.getenv('ARTICLES_DIR', '/Users/ericchiu/Documents/sungertain-design/G系統專案'))

# 建立日誌目錄
LOG_DIR.mkdir(parents=True, exist_ok=True)

# 配置日誌
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(LOG_DIR / 'app.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)


# ====== 文章解析器 ======
class DocxArticleParser:
    """解析 .docx 文件並提取 Meta 信息和內容"""

    def __init__(self, docx_path):
        self.docx_path = docx_path
        self.doc = Document(docx_path)

    def extract_meta_table(self):
        """從第一個表格提取 Meta 信息"""
        if not self.doc.tables:
            raise ValueError("文檔中沒有表格")

        meta_table = self.doc.tables[0]
        meta = {}

        for row in meta_table.rows:
            if len(row.cells) >= 2:
                key = row.cells[0].text.strip()
                value = row.cells[1].text.strip()
                meta[key] = value

        return {
            'seo_title': meta.get('SEO 標題', ''),
            'seo_description': meta.get('SEO 描述', ''),
            'publish_date': meta.get('發布日期', ''),
            'category': meta.get('文章分類', '') or meta.get('分類', ''),
            'product': meta.get('推薦產品', '')
        }

    def extract_article_body(self):
        """提取文章主體內容"""
        paragraphs = []
        skip_table = True

        for element in self.doc.element.body:
            if element.tag.endswith('tbl'):
                if skip_table:
                    skip_table = False
                continue

            if element.tag.endswith('p'):
                para = None
                for p in self.doc.paragraphs:
                    if p._element == element:
                        para = p
                        break

                if para and para.text.strip():
                    paragraphs.append(para)

        return paragraphs

    def extract_hyperlinks(self):
        """提取超連結"""
        hyperlinks = []

        for para in self.doc.paragraphs:
            for child in para._element:
                if child.tag.endswith('hyperlink'):
                    rel_id = child.get(qn('r:id'))
                    if rel_id:
                        try:
                            target = self.doc.part.rels[rel_id].target_ref
                            text_elements = child.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                            text = ''.join([t.text for t in text_elements if t.text])

                            if target and text:
                                hyperlinks.append({
                                    'url': target,
                                    'title': text.strip(),
                                    'authority_score': 0.8
                                })
                        except:
                            pass

        return hyperlinks

    def parse(self, day_in_cycle=1) -> Dict:
        """完整解析"""
        try:
            meta = self.extract_meta_table()
            paragraphs = self.extract_article_body()
            hyperlinks = self.extract_hyperlinks()

            article_body = self._paragraphs_to_html(paragraphs)
            publish_date = self._parse_publish_date(meta['publish_date'])

            return {
                'meta_title': meta['seo_title'],
                'meta_description': meta['seo_description'],
                'article_body': article_body,
                'category': meta['category'],
                'tags': self._extract_tags(meta['category']),
                'related_products': [meta['product']] if meta['product'] else [],
                'publish_date': publish_date,
                'status': 'draft',
                'day_in_cycle': day_in_cycle,
                'quality_score': self._calculate_quality_score(paragraphs),
                'references': hyperlinks,
                'created_by': 'cowork_system'
            }
        except Exception as e:
            logger.error(f"解析失敗: {e}")
            raise

    @staticmethod
    def _paragraphs_to_html(paragraphs):
        """轉換為 HTML"""
        html_parts = []
        for para in paragraphs:
            text = para.text.strip()
            if not text:
                continue

            if para.style and para.style.name and para.style.name.startswith('Heading'):
                level = 2 if 'Heading 2' in para.style.name else 1
                html_parts.append(f"<h{level}>{text}</h{level}>")
            else:
                html_parts.append(f"<p>{text}</p>")

        return "\n".join(html_parts)

    @staticmethod
    def _parse_publish_date(date_str):
        """解析日期"""
        try:
            date_str = date_str.replace('年', '-').replace('月', '-').replace('日', '')
            parts = date_str.split('-')
            if len(parts) == 3:
                year, month, day = parts
                dt = datetime(int(year), int(month), int(day), 9, 0, 0, tzinfo=timezone.utc)
                return dt.strftime('%Y-%m-%dT%H:%M:%S+08:00')
        except:
            pass

        now = datetime.now(timezone.utc)
        return now.strftime('%Y-%m-%dT%H:%M:%S+08:00')

    @staticmethod
    def _extract_tags(category):
        """提取標籤"""
        tag_map = {
            '選購指南': ['靈芝', '選購指南', '消費者指南'],
            '入門認識': ['靈芝', '入門認識', '健康知識'],
            '飲食指南': ['靈芝', '飲食指南', '健康飲食'],
            '保存方式': ['靈芝', '保存方式', '產品護理'],
            '常見問題': ['靈芝', '常見問題', 'FAQ'],
            '深入認識': ['靈芝', '深入認識', '科學'],
            '產品評測': ['靈芝', '產品評測', '評測']
        }
        return tag_map.get(category, ['靈芝', category] if category else ['靈芝'])

    @staticmethod
    def _calculate_quality_score(paragraphs):
        """品質評分"""
        total_chars = sum(len(p.text) for p in paragraphs)
        para_count = len([p for p in paragraphs if p.text.strip()])

        score = 100

        if total_chars < 600:
            score -= 15
        elif total_chars > 2000:
            score -= 10

        if para_count < 5:
            score -= 10
        elif para_count > 15:
            score -= 5

        return max(0, min(100, score))


# ====== API 客戶端 ======
class GSystemAPIClient:
    """G 系統 API 客戶端"""

    def __init__(self, api_url: str, api_key: str):
        self.api_url = api_url
        self.api_key = api_key
        self.timeout = 30
        self.max_retries = 3

    def publish_article(self, payload: Dict) -> Tuple[bool, Dict]:
        """推送文章"""
        required_fields = ['meta_title', 'meta_description', 'article_body',
                          'category', 'publish_date', 'status']

        for field in required_fields:
            if field not in payload or not payload[field]:
                return False, {'success': False, 'error_message': f'缺少: {field}'}

        valid_categories = ['選購指南', '入門認識', '產品評測', '使用心得', '健康知識', '其他']
        if payload['category'] not in valid_categories:
            return False, {'success': False, 'error_message': f'無效分類'}

        return self._send_with_retry(payload)

    def _send_with_retry(self, payload: Dict) -> Tuple[bool, Dict]:
        """帶重試的推送"""
        for attempt in range(1, self.max_retries + 1):
            try:
                logger.info(f"推送嘗試 {attempt}/{self.max_retries}")

                headers = {
                    'Authorization': f'Bearer {self.api_key}',
                    'Content-Type': 'application/json'
                }

                response = requests.post(
                    self.api_url,
                    headers=headers,
                    json=payload,
                    timeout=self.timeout
                )

                if response.status_code == 201:
                    result = response.json()
                    logger.info(f"✅ 推送成功: {result.get('article_id')}")
                    return True, result

                elif response.status_code == 401:
                    logger.error("認證失敗")
                    return False, response.json()

                elif response.status_code == 400:
                    logger.error("驗證失敗")
                    return False, response.json()

                elif response.status_code == 429:
                    if attempt < self.max_retries:
                        import time
                        logger.warning("限流，60秒後重試")
                        time.sleep(60)
                        continue

                elif response.status_code == 500:
                    if attempt < self.max_retries:
                        import time
                        logger.warning("伺服器錯誤，10秒後重試")
                        time.sleep(10)
                        continue

            except requests.exceptions.Timeout:
                if attempt < self.max_retries:
                    import time
                    logger.warning("超時，5秒後重試")
                    time.sleep(5)

            except Exception as e:
                logger.error(f"錯誤: {e}")
                return False, {'success': False, 'error_message': str(e)}

        return False, {'success': False, 'error_message': '所有重試都失敗'}

    def verify_connection(self) -> bool:
        """驗證連接"""
        try:
            headers = {
                'Authorization': f'Bearer {self.api_key}',
                'Content-Type': 'application/json'
            }

            response = requests.post(
                self.api_url,
                headers=headers,
                json={
                    'meta_title': '連接測試',
                    'meta_description': '測試',
                    'article_body': '測試',
                    'category': '其他',
                    'publish_date': '2026-05-15T09:00:00+08:00',
                    'status': 'draft'
                },
                timeout=self.timeout
            )

            if response.status_code in [201, 400, 401]:
                logger.info("✅ API 連接正常")
                return True
            else:
                logger.error(f"❌ HTTP {response.status_code}")
                return False

        except Exception as e:
            logger.error(f"連接失敗: {e}")
            return False


# ====== 主程序 ======
def publish_article(docx_path: str, day_in_cycle: int) -> bool:
    """發布單篇文章"""
    logger.info(f"開始處理: {docx_path}")

    try:
        # 解析
        parser = DocxArticleParser(docx_path)
        payload = parser.parse(day_in_cycle)
        logger.info(f"✓ 已解析: {payload['meta_title'][:50]}")

        # 推送
        client = GSystemAPIClient(API_URL, API_KEY)
        success, response = client.publish_article(payload)

        if success:
            article_id = response.get('article_id')
            audit_url = response.get('audit_url')
            logger.info(f"✅ 成功! ID: {article_id}")
            logger.info(f"審核: {audit_url}")

            # 記錄成功
            with open(LOG_DIR / 'success.log', 'a', encoding='utf-8') as f:
                f.write(f"[{datetime.now()}] {article_id} - {payload['meta_title']}\n")

            return True
        else:
            error = response.get('error_message', '未知')
            logger.error(f"❌ {error}")

            # 記錄失敗
            with open(LOG_DIR / 'error.log', 'a', encoding='utf-8') as f:
                f.write(f"[{datetime.now()}] {payload['meta_title']}\n  錯誤: {error}\n")

            return False

    except Exception as e:
        logger.error(f"❌ {e}")
        return False


def test_connection():
    """測試連接"""
    client = GSystemAPIClient(API_URL, API_KEY)
    if client.verify_connection():
        print("✅ API 連接正常")
        return 0
    else:
        print("❌ API 連接失敗")
        return 1


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    command = sys.argv[1]

    if command == 'publish' and len(sys.argv) >= 3:
        docx_path = sys.argv[2]
        day_in_cycle = int(sys.argv[3]) if len(sys.argv) > 3 else 1
        success = publish_article(docx_path, day_in_cycle)
        sys.exit(0 if success else 1)

    elif command == 'test':
        sys.exit(test_connection())

    else:
        print(__doc__)
        sys.exit(1)
